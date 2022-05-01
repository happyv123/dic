import docx
import time
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from progress.bar import IncrementalBar
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By

options = Options()
options.add_argument('--headless')
options.add_argument('--disable-logging')
options.add_argument('--disable-extensions')
options.add_argument('--log-level=3')
options.add_argument('--disable-in-process-stack-traces')
options.add_argument('--ignore-certificate-errors')

s = Service('chromedriver.exe')
driver = webdriver.Chrome(options=options, service=s)

URL_TEMPLATE = "https://dictionary.cambridge.org/ru/%D1%81%D0%BB%D0%BE%D0%B2%D0%B0%D1%80%D1%8C/%D0%B0%D0%BD%D0%B3%D0%BB%D0%BE-%D1%80%D1%83%D1%81%D1%81%D0%BA%D0%B8%D0%B9/"

mydoc = docx.Document()
style = mydoc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

with open('wordlist.txt', 'r') as f:
    wordlist = f.readlines()
    bar = IncrementalBar('Process', max=70)
    success_word_counter: int = 1
    all_word_counter = 0
    while True:
        try:
            word = wordlist[all_word_counter]
            word = word.replace("\n", "")
            driver.get(URL_TEMPLATE + word)
        except:
            print("No more words")
            break

        try:
            word_def = driver.find_element(By.XPATH, value="//div[@class='def ddef_d db']")
            current_par = mydoc.add_paragraph()
            bold_word = current_par.add_run(str(success_word_counter) + "." + word)
            bold_word.bold = True
            italic_def = current_par.add_run(' - ' + word_def.text).add_break(WD_BREAK.LINE)
            success_word_counter += 1
            bar.next()
        except:
            print('\n Error. No such word: ' + word)
            all_word_counter += 1
            continue

        try:
            example = driver.find_element(By.XPATH, value="//div[@class='examp dexamp']")
            italic_example = current_par.add_run(example.text)
            italic_example.italic = True
        except:
            print('\n Error. No example for word: ' + word)

        all_word_counter += 1
        mydoc.save('dic.docx')
        time.sleep(1)
        if success_word_counter == 70:
            break
            
bar.finish()
driver.quit()

print("Check the docx file for errors!!!")
